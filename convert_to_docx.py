from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

def create_project_report():
    """Create a Word document for the project report"""
    
    # Create document
    doc = Document()
    
    # Set up styles
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Title Page
    title = doc.add_heading('PREDICTIVE MAINTENANCE DASHBOARD FOR SOLAR PANELS', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Major Project Report')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(16)
    subtitle.runs[0].font.bold = True
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    # Add image if available
    screenshot_path = r"C:\Users\Imamsab jamdar\.gemini\antigravity\brain\191c3fa0-c7a0-4d2a-b173-bc27da7a4fc1\dashboard_main_view_1764654012628.png"
    if os.path.exists(screenshot_path):
        doc.add_picture(screenshot_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # Executive Summary
    doc.add_heading('EXECUTIVE SUMMARY', 1)
    doc.add_paragraph(
        'This project presents a comprehensive Predictive Maintenance Dashboard for solar panel systems '
        'that leverages machine learning algorithms to predict maintenance needs, optimize performance, '
        'and minimize downtime. The system combines real-time monitoring, historical data analysis, and '
        'predictive analytics to provide actionable insights for solar panel maintenance and performance optimization.'
    )
    doc.add_paragraph(
        'The application is built using modern web technologies with a React-based frontend and a Python '
        'FastAPI backend, incorporating machine learning models for predictive analytics. The system monitors '
        'key performance indicators (KPIs), analyzes trends, and provides maintenance recommendations based on '
        'real-time and historical data.'
    )
    
    # 1. Project Overview
    doc.add_heading('1. PROJECT OVERVIEW', 1)
    
    doc.add_heading('1.1 Project Title', 2)
    doc.add_paragraph('Predictive Maintenance Dashboard for Solar Panel Systems')
    
    doc.add_heading('1.2 Objective', 2)
    doc.add_paragraph('To develop an intelligent monitoring and predictive maintenance system that:')
    objectives = [
        'Monitors solar panel performance in real-time',
        'Predicts maintenance requirements using machine learning',
        'Optimizes energy output through proactive maintenance',
        'Reduces operational costs and system downtime',
        'Provides data-driven insights for decision-making'
    ]
    for obj in objectives:
        doc.add_paragraph(obj, style='List Bullet')
    
    doc.add_heading('1.3 Scope', 2)
    doc.add_paragraph('The system encompasses:')
    scope_items = [
        'Real-time data collection and monitoring',
        'Historical data analysis and visualization',
        'Machine learning-based predictions',
        'Maintenance task management',
        'Performance analytics and reporting',
        'User-friendly web interface'
    ]
    for item in scope_items:
        doc.add_paragraph(item, style='List Bullet')
    
    # 2. System Architecture
    doc.add_heading('2. SYSTEM ARCHITECTURE', 1)
    
    doc.add_heading('2.1 Technology Stack', 2)
    
    doc.add_heading('Frontend Technologies', 3)
    frontend_tech = [
        'Framework: React 19.0.0',
        'Routing: React Router DOM 7.5.1',
        'UI Components: Radix UI component library',
        'Styling: Tailwind CSS 3.4.17',
        'Charts: Recharts 3.1.2',
        'HTTP Client: Axios 1.8.4',
        'Form Management: React Hook Form 7.56.2',
        'Build Tool: CRACO (Create React App Configuration Override)'
    ]
    for tech in frontend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('Backend Technologies', 3)
    backend_tech = [
        'Framework: FastAPI 0.110.1',
        'Server: Uvicorn 0.25.0',
        'Database: MongoDB (Motor async driver 3.3.1)',
        'Machine Learning: scikit-learn ≥1.3.0, NumPy ≥1.26.0, Pandas ≥2.2.0',
        'Model Persistence: Joblib ≥1.3.0',
        'Environment Management: python-dotenv ≥1.0.1'
    ]
    for tech in backend_tech:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_heading('Machine Learning Models', 3)
    ml_models = [
        'Efficiency Prediction Model: Random Forest Regressor',
        'Power Output Prediction Model: Random Forest Regressor',
        'Maintenance Classification Model: Random Forest Classifier',
        'Feature Scaling: StandardScaler'
    ]
    for model in ml_models:
        doc.add_paragraph(model, style='List Bullet')
    
    # 3. Features and Functionality
    doc.add_heading('3. FEATURES AND FUNCTIONALITY', 1)
    
    doc.add_heading('3.1 Dashboard Module', 2)
    doc.add_paragraph('The main dashboard provides a comprehensive overview of the solar panel system:')
    
    doc.add_heading('Key Performance Indicators (KPIs)', 3)
    kpis = [
        'Current Power Output (kW): Real-time power generation monitoring with trend indicators',
        'System Efficiency (%): Overall system efficiency calculation with historical comparisons',
        'Panel Temperature (°C): Real-time temperature monitoring with optimal range indicators (25-35°C)',
        'Dust Level Index: Categorized as Low/Medium/High with color-coded status indicators'
    ]
    for kpi in kpis:
        doc.add_paragraph(kpi, style='List Bullet')
    
    doc.add_heading('Visualization Components', 3)
    viz_components = [
        'Power Output Trend Chart: 24-hour power generation curve',
        'Efficiency Analysis Chart: System efficiency over time',
        'Efficiency Distribution: Pie chart showing performance breakdown',
        'Maintenance Status Panel: Overview of panel health',
        'Recent Alerts: Real-time system notifications'
    ]
    for comp in viz_components:
        doc.add_paragraph(comp, style='List Bullet')
    
    doc.add_heading('3.2 Performance Module', 2)
    doc.add_paragraph('Detailed performance analytics and historical data visualization:')
    perf_features = [
        'Multi-day performance analysis (1, 7, 30, 90 days)',
        'Power output vs. expected power comparison',
        'Efficiency trends and patterns',
        'Environmental factor correlation (temperature, humidity, irradiance)',
        'Performance deviation analysis',
        'Downloadable performance reports (PDF)'
    ]
    for feature in perf_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('3.3 Maintenance Module', 2)
    doc.add_paragraph('Comprehensive maintenance management system:')
    maint_features = [
        'Maintenance task scheduling',
        'Task priority management (Low, Medium, High, Critical)',
        'Task status tracking (Scheduled, In Progress, Completed, Cancelled)',
        'Assigned technician management',
        'Task filtering and search',
        'Automated maintenance recommendations',
        'Task completion tracking'
    ]
    for feature in maint_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('3.4 Predictions Module', 2)
    doc.add_paragraph('Machine learning-powered predictive analytics:')
    
    pred_types = [
        'Power Output Forecast (7 Days): Predicted power generation with confidence intervals',
        'Efficiency Forecast: Future efficiency trends and degradation predictions',
        'Maintenance Predictions: Predicted maintenance requirements with risk probability',
        'AI-Driven Recommendations: Data-driven maintenance suggestions and optimization tips'
    ]
    for pred in pred_types:
        doc.add_paragraph(pred, style='List Bullet')
    
    # Add predictions screenshot if available
    predictions_path = r"C:\Users\Imamsab jamdar\.gemini\antigravity\brain\191c3fa0-c7a0-4d2a-b173-bc27da7a4fc1\predictions_page_view_1764654066655.png"
    if os.path.exists(predictions_path):
        doc.add_paragraph()
        doc.add_picture(predictions_path, width=Inches(5.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption = doc.add_paragraph('Figure: Predictions Dashboard')
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption.runs[0].font.italic = True
    
    doc.add_heading('3.5 Settings Module', 2)
    settings_features = [
        'User management (Create, Read, Update, Delete)',
        'Role-based access control (Admin, Technician, Viewer)',
        'System threshold configuration',
        'Alert settings customization',
        'Data import/export functionality',
        'CSV data upload for model training',
        'System preferences'
    ]
    for feature in settings_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # 4. Machine Learning Implementation
    doc.add_heading('4. MACHINE LEARNING IMPLEMENTATION', 1)
    
    doc.add_heading('4.1 Data Processing Pipeline', 2)
    doc.add_paragraph('Feature Engineering:')
    features = [
        'Irradiance (W/m²)',
        'Temperature (°C)',
        'Humidity (%)',
        'Dust Level (%)',
        'Voltage (V)',
        'Current (A)',
        'Hour of day',
        'Day of year',
        'Month'
    ]
    for feature in features:
        doc.add_paragraph(feature, style='List Bullet')
    
    doc.add_heading('4.2 Model Training Process', 2)
    training_steps = [
        'Data Collection: Historical solar panel data from MongoDB',
        'Feature Extraction: Time-based and environmental features',
        'Data Preprocessing: StandardScaler normalization',
        'Model Training: Random Forest algorithms (100 estimators)',
        'Model Evaluation: Train-test split (80-20)',
        'Model Persistence: Saved as .pkl files using joblib'
    ]
    for step in training_steps:
        doc.add_paragraph(step, style='List Bullet')
    
    doc.add_heading('4.3 Prediction Models', 2)
    
    doc.add_paragraph('Efficiency Prediction Model:', style='List Bullet')
    doc.add_paragraph('Algorithm: Random Forest Regressor')
    doc.add_paragraph('Purpose: Predict future panel efficiency')
    doc.add_paragraph('Output: Efficiency percentage (0-100%)')
    doc.add_paragraph()
    
    doc.add_paragraph('Power Output Prediction Model:', style='List Bullet')
    doc.add_paragraph('Algorithm: Random Forest Regressor')
    doc.add_paragraph('Purpose: Forecast power generation')
    doc.add_paragraph('Output: Power output in kW')
    doc.add_paragraph()
    
    doc.add_paragraph('Maintenance Classification Model:', style='List Bullet')
    doc.add_paragraph('Algorithm: Random Forest Classifier')
    doc.add_paragraph('Purpose: Classify maintenance requirements')
    doc.add_paragraph('Classes: Not Required (0), Warning (1), Required (2)')
    
    # 5. Database Schema
    doc.add_heading('5. DATABASE SCHEMA', 1)
    
    doc.add_heading('5.1 Collections', 2)
    
    doc.add_paragraph('solar_data Collection:', style='Heading 3')
    solar_fields = [
        'id: String (UUID)',
        'timestamp: DateTime',
        'power_output: Float',
        'irradiance: Float',
        'temperature: Float',
        'humidity: Float',
        'dust_level: Float',
        'voltage: Float',
        'current: Float',
        'efficiency: Float',
        'maintenance_status: String',
        'maintenance_description: String'
    ]
    for field in solar_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_paragraph('users Collection:', style='Heading 3')
    user_fields = [
        'id: String (UUID)',
        'name: String',
        'email: String',
        'role: String',
        'status: String',
        'created_at: DateTime'
    ]
    for field in user_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    doc.add_paragraph('maintenance_tasks Collection:', style='Heading 3')
    task_fields = [
        'id: String (UUID)',
        'panel_id: String',
        'task_type: String',
        'description: String',
        'priority: String',
        'scheduled_date: DateTime',
        'estimated_duration: String',
        'assigned_to: String',
        'status: String',
        'created_at: DateTime',
        'completed_at: DateTime'
    ]
    for field in task_fields:
        doc.add_paragraph(field, style='List Bullet')
    
    # 6. API Endpoints
    doc.add_heading('6. API ENDPOINTS', 1)
    
    doc.add_heading('6.1 Dashboard Endpoints', 2)
    dashboard_endpoints = [
        'GET /api/dashboard/kpis - Retrieve current KPIs',
        'GET /api/dashboard/performance?days={n} - Get performance data for n days',
        'GET /api/dashboard/alerts - Fetch active maintenance alerts'
    ]
    for endpoint in dashboard_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')
    
    doc.add_heading('6.2 Predictions Endpoints', 2)
    doc.add_paragraph('GET /api/predictions/forecast - Get 7-day ML predictions', style='List Bullet')
    
    doc.add_heading('6.3 Maintenance Endpoints', 2)
    maint_endpoints = [
        'GET /api/maintenance/tasks - List all maintenance tasks',
        'POST /api/maintenance/tasks - Create new maintenance task',
        'PUT /api/maintenance/tasks/{id} - Update task status',
        'DELETE /api/maintenance/tasks/{id} - Delete maintenance task'
    ]
    for endpoint in maint_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')
    
    doc.add_heading('6.4 Data Management Endpoints', 2)
    data_endpoints = [
        'POST /api/data/import - Import CSV data',
        'POST /api/data/train - Trigger model training',
        'GET /api/data/export - Export data as CSV'
    ]
    for endpoint in data_endpoints:
        doc.add_paragraph(endpoint, style='List Bullet')
    
    # 7. Installation and Deployment
    doc.add_heading('7. INSTALLATION AND DEPLOYMENT', 1)
    
    doc.add_heading('7.1 Prerequisites', 2)
    prereqs = [
        'Python: 3.8 or higher',
        'Node.js: 14.x or higher',
        'MongoDB: 4.4 or higher',
        'npm/yarn: Latest version'
    ]
    for prereq in prereqs:
        doc.add_paragraph(prereq, style='List Bullet')
    
    doc.add_heading('7.2 Backend Setup', 2)
    doc.add_paragraph('Step 1: Navigate to backend directory')
    code_para = doc.add_paragraph('cd backend')
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph('Step 2: Install Python dependencies')
    code_para = doc.add_paragraph('pip install -r requirements.txt')
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph('Step 3: Start the backend server')
    code_para = doc.add_paragraph('python -m uvicorn server:app --reload --host 0.0.0.0 --port 8000')
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(10)
    
    doc.add_heading('7.3 Frontend Setup', 2)
    doc.add_paragraph('Step 1: Navigate to frontend directory')
    code_para = doc.add_paragraph('cd frontend')
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph('Step 2: Install Node.js dependencies')
    code_para = doc.add_paragraph('npm install')
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(10)
    doc.add_paragraph()
    doc.add_paragraph('Step 3: Start the development server')
    code_para = doc.add_paragraph('npm start')
    code_para.runs[0].font.name = 'Courier New'
    code_para.runs[0].font.size = Pt(10)
    
    doc.add_heading('7.4 Access Points', 2)
    access_points = [
        'Frontend Application: http://localhost:3000',
        'Backend API: http://localhost:8000',
        'API Documentation: http://localhost:8000/docs',
        'ReDoc: http://localhost:8000/redoc'
    ]
    for point in access_points:
        doc.add_paragraph(point, style='List Bullet')
    
    # 8. Key Benefits
    doc.add_heading('8. KEY BENEFITS', 1)
    
    doc.add_heading('8.1 Operational Benefits', 2)
    op_benefits = [
        'Reduced Downtime: Predictive maintenance prevents unexpected failures',
        'Cost Optimization: Proactive maintenance reduces emergency repair costs',
        'Extended Equipment Life: Timely maintenance prolongs panel lifespan',
        'Improved Efficiency: Optimized cleaning and maintenance schedules',
        'Data-Driven Decisions: Analytics-based operational strategies'
    ]
    for benefit in op_benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_heading('8.2 Technical Benefits', 2)
    tech_benefits = [
        'Real-time Monitoring: Instant visibility into system performance',
        'Automated Alerts: Proactive notification of potential issues',
        'Scalability: MongoDB and FastAPI support large-scale deployments',
        'Accuracy: Machine learning improves prediction accuracy over time',
        'Integration Ready: RESTful API for third-party integrations'
    ]
    for benefit in tech_benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    doc.add_heading('8.3 Business Benefits', 2)
    biz_benefits = [
        'ROI Improvement: Maximized energy output and reduced costs',
        'Compliance: Automated maintenance logs for regulatory requirements',
        'Reporting: Comprehensive performance reports for stakeholders',
        'Competitive Advantage: Advanced analytics capabilities',
        'Sustainability: Optimized renewable energy production'
    ]
    for benefit in biz_benefits:
        doc.add_paragraph(benefit, style='List Bullet')
    
    # 9. Future Enhancements
    doc.add_heading('9. FUTURE ENHANCEMENTS', 1)
    
    doc.add_heading('9.1 Planned Features', 2)
    future_features = [
        'Mobile Application: iOS and Android apps for remote monitoring',
        'Advanced Analytics: Deep learning models for improved predictions',
        'IoT Integration: Direct sensor data integration',
        'Weather API Integration: Real-time weather data for better predictions',
        'Multi-site Management: Support for multiple solar installations',
        'Automated Reporting: Scheduled email reports',
        'Anomaly Detection: AI-powered anomaly identification',
        'Energy Storage Integration: Battery system monitoring',
        'Cost Analysis Module: Financial impact analysis',
        'Predictive Alerts: SMS/Email notifications'
    ]
    for feature in future_features:
        doc.add_paragraph(feature, style='List Bullet')
    
    # 10. Conclusion
    doc.add_heading('10. CONCLUSION', 1)
    doc.add_paragraph(
        'The Predictive Maintenance Dashboard for Solar Panels represents a comprehensive solution for modern '
        'solar energy management. By combining real-time monitoring, historical data analysis, and machine learning '
        'predictions, the system enables proactive maintenance strategies that optimize performance, reduce costs, '
        'and extend equipment lifespan.'
    )
    doc.add_paragraph(
        'The application demonstrates the effective integration of modern web technologies, machine learning algorithms, '
        'and database systems to create a production-ready solution for renewable energy management. The modular '
        'architecture ensures scalability and maintainability, while the user-friendly interface makes advanced '
        'analytics accessible to all stakeholders.'
    )
    doc.add_paragraph(
        'This project showcases the practical application of full-stack web development, machine learning in production '
        'environments, real-time data processing and visualization, RESTful API design, modern UI/UX principles, and '
        'database design and optimization. The system is ready for deployment and can be extended with additional '
        'features to meet evolving business requirements.'
    )
    
    # Appendix
    doc.add_page_break()
    doc.add_heading('APPENDIX', 1)
    
    doc.add_heading('A. System Requirements', 2)
    doc.add_paragraph('Minimum Hardware Requirements:')
    min_req = [
        'CPU: Dual-core processor, 2.0 GHz',
        'RAM: 4 GB',
        'Storage: 10 GB available space',
        'Network: Broadband internet connection'
    ]
    for req in min_req:
        doc.add_paragraph(req, style='List Bullet')
    
    doc.add_paragraph('Recommended Hardware Requirements:')
    rec_req = [
        'CPU: Quad-core processor, 3.0 GHz',
        'RAM: 8 GB or higher',
        'Storage: 20 GB SSD',
        'Network: High-speed internet connection'
    ]
    for req in rec_req:
        doc.add_paragraph(req, style='List Bullet')
    
    # Footer
    doc.add_page_break()
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Project Developed By: [Your Name]\n')
    footer_run.font.size = Pt(12)
    footer_run = footer_para.add_run('Institution: [Your Institution]\n')
    footer_run.font.size = Pt(12)
    footer_run = footer_para.add_run('Date: December 2, 2025\n')
    footer_run.font.size = Pt(12)
    footer_run = footer_para.add_run('Version: 1.0')
    footer_run.font.size = Pt(12)
    
    # Save document
    output_path = r"C:\Users\Imamsab jamdar\OneDrive\Desktop\Predictive_maintenances-main\Predictive_Maintenance_Dashboard_Project_Report.docx"
    doc.save(output_path)
    print(f"✅ Project report successfully created: {output_path}")
    return output_path

if __name__ == "__main__":
    create_project_report()
